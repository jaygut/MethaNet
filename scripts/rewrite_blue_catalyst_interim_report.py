#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def build_report() -> Path:
    root = Path("/Users/jaygut/Documents/Side_Projects/MethaNet")
    snapshot = (
        root
        / "results/blue_catalyst_poc/interim_snapshots"
        / "apolo_full_20260228_080644_embed_20260305_061952_notebook_interim_20260305_220228"
    )

    summary_path = snapshot / "analytics/analytics_summary.json"
    fig_dir = snapshot / "figures"
    table_dir = snapshot / "tables"
    out_doc = (
        root
        / "results/blue_catalyst_poc/runs/apolo_20260226_194505/report/blue_catalyst_poc_report.docx"
    )

    summary = json.loads(summary_path.read_text())
    alpha = pd.read_csv(table_dir / "alpha_transfer_candidates_top40.tsv", sep="\t")
    bridge = pd.read_csv(table_dir / "bridge_genomes_top30_partial.tsv", sep="\t")

    fig_order = [
        (
            "projection_landscapes_panel.png",
            "Figure 1. Projection landscapes (PCA, UMAP, t-SNE) with neighborhood entropy distribution.",
        ),
        (
            "pca_scree.png",
            "Figure 2. PCA variance spectrum in the embedding space.",
        ),
        (
            "signal_qc_panel.png",
            "Figure 3. Signal QC panel: proteins-per-genome, proteome-size distribution, mixing z-scores, and nearest-neighbor geometry.",
        ),
        (
            "metric_rigor_panel.png",
            "Figure 4. Statistical rigor panel with silhouette bootstrap and consolidated separation metrics.",
        ),
        (
            "silhouette_by_ecosystem.png",
            "Figure 5. Per-genome silhouette distribution by ecosystem.",
        ),
        (
            "ecosystem_trajectory_axis.png",
            "Figure 6. Projection densities along the rumen-to-wetland centroid axis.",
        ),
        (
            "cosine_distance_heatmap_sampled.png",
            "Figure 7. Sampled pairwise cosine distance heatmap showing global structure.",
        ),
    ]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Blue Catalyst Interim Embedding Analytics Report")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Full rewrite focused exclusively on the interim snapshot").italic = True

    p_snapshot = doc.add_paragraph()
    p_snapshot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_snapshot.add_run(f"Snapshot: {summary['snapshot_name']}")

    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_src.add_run(f"Source artifacts: {summary['source_artifacts_dir']}")

    doc.add_page_break()

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This document replaces the prior Blue Catalyst report with an interim-run-only narrative. "
        "All statistics and figures are taken from the specified interim snapshot."
    )

    sil = summary["silhouette_global"]
    bs = summary["silhouette_bootstrap"]
    perm = summary["permanova"]
    cls = summary["classifier_cv"]
    conf = summary["confounding_assessment"]

    metrics = [
        f"Embeddings analyzed: {summary['n_embeddings']} genomes (rumen={summary['ecosystem_counts'].get('rumen', 0)}, wetland={summary['ecosystem_counts'].get('wetland', 0)}).",
        f"Embedding dimension: {summary['embedding_dim']}.",
        f"Global silhouette: {sil:.4f}; bootstrap mean={bs['mean']:.4f}; 95% CI=[{bs['ci95_low']:.4f}, {bs['ci95_high']:.4f}] (n_boot={bs['n_boot']}).",
        f"PERMANOVA: pseudo-F={perm['pseudo_f']:.3f}, R^2={perm['r2']:.4f}, p={perm['p_value']:.4g} ({perm['permutations']} permutations).",
        f"CV classifier: AUC={cls['auc']:.4f}, AUPRC={cls['auprc']:.4f}, balanced accuracy={cls['balanced_accuracy']:.4f}.",
        f"Source purity by ecosystem: {conf.get('source_purity_by_ecosystem', {})}.",
        f"Unknown domain fraction by ecosystem: {conf.get('domain_unknown_fraction_by_ecosystem', {})}.",
    ]
    for line in metrics:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. Scientific Caveats", level=1)
    for caveat in summary.get("scientific_caveats", []):
        doc.add_paragraph(caveat, style="List Bullet")
    doc.add_paragraph(
        "Interpretation policy: this interim run captures high-contrast cohort structure, "
        "but not fully deconfounded ecosystem biology."
    )

    doc.add_heading("3. Methods (Interim Snapshot Scope)", level=1)
    methods = [
        "Input matrix: checkpoint-aggregated genome embeddings with per-genome metadata.",
        "Distance metric: cosine distance in the latent embedding space.",
        "Projection methods: PCA, UMAP, and t-SNE.",
        "Neighborhood features: opposite-neighbor fraction, bridge entropy, expected mixing baseline, and mixing z-score.",
        "Boundary features: nearest same-ecosystem and nearest opposite-ecosystem distances.",
        "Alpha transfer score: mean(z(bridge_entropy), z(boundary_balance), z(-local_mean_cosine_distance)).",
        "Rigor checks: silhouette bootstrap confidence intervals and PERMANOVA permutation testing.",
    ]
    for line in methods:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("4. Figure-Centered Results", level=1)
    for fname, caption in fig_order:
        fig_path = fig_dir / fname
        if not fig_path.exists():
            continue
        doc.add_picture(str(fig_path), width=Inches(6.8))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        doc.add_paragraph("")

    doc.add_heading("5. Alpha Candidate Prioritization", level=1)
    doc.add_paragraph(
        "Candidates are prioritized for follow-up based on boundary proximity and cross-ecosystem mixing signal in embedding space."
    )

    alpha_cols = [
        "sample",
        "ecosystem",
        "source",
        "alpha_transfer_score",
        "bridge_entropy",
        "mixing_zscore",
    ]
    doc.add_heading("5.1 Top 10 alpha-transfer candidates", level=3)
    t = doc.add_table(rows=1, cols=len(alpha_cols))
    t.style = "Light List Accent 1"
    for i, c in enumerate(alpha_cols):
        t.rows[0].cells[i].text = c
    for _, row in alpha.sort_values("alpha_transfer_score", ascending=False).head(10).iterrows():
        cells = t.add_row().cells
        for i, c in enumerate(alpha_cols):
            v = row[c]
            cells[i].text = f"{v:.4f}" if isinstance(v, float) else str(v)

    bridge_cols = [
        "sample",
        "ecosystem",
        "source",
        "bridge_entropy",
        "mixing_zscore",
        "nearest_opposite_distance",
    ]
    doc.add_heading("5.2 Top 10 bridge-entropy candidates", level=3)
    t2 = doc.add_table(rows=1, cols=len(bridge_cols))
    t2.style = "Light List Accent 1"
    for i, c in enumerate(bridge_cols):
        t2.rows[0].cells[i].text = c
    for _, row in bridge.sort_values("bridge_entropy", ascending=False).head(10).iterrows():
        cells = t2.add_row().cells
        for i, c in enumerate(bridge_cols):
            v = row[c]
            cells[i].text = f"{v:.6f}" if isinstance(v, float) else str(v)

    doc.add_heading("6. Interim Conclusions", level=1)
    conclusions = [
        "Embedding-space separation is robust in this interim run.",
        "Confounding between ecosystem and source remains strong and must be explicitly controlled in follow-up analyses.",
        "The alpha-transfer ranking provides a practical shortlist for marker-aware and taxonomic validation.",
        "This report is an interim decision-support artifact, not a final ecological causality claim.",
    ]
    for line in conclusions:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("7. Artifact Index", level=1)
    idx_lines = [
        f"analytics_summary.json: {summary_path}",
        f"alpha_transfer_candidates_top40.tsv: {table_dir / 'alpha_transfer_candidates_top40.tsv'}",
        f"alpha_transfer_top20_per_ecosystem.tsv: {table_dir / 'alpha_transfer_top20_per_ecosystem.tsv'}",
        f"bridge_genomes_top30_partial.tsv: {table_dir / 'bridge_genomes_top30_partial.tsv'}",
        f"embedding_projection_partial.tsv: {table_dir / 'embedding_projection_partial.tsv'}",
    ]
    for line in idx_lines:
        doc.add_paragraph(line, style="List Bullet")

    out_doc.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_doc)
    return out_doc


if __name__ == "__main__":
    out = build_report()
    print(f"Wrote: {out}")
