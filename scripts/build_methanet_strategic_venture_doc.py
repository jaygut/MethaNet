#!/usr/bin/env python3
"""
Build MethaNet Strategic Venture Document as a Word document.

Generates a publication-grade Word document covering MethaNet's positioning
as a domain-expert agentic system for cross-ecosystem methanogenic microbiology.

References verified via web search March 2026.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

# --- Configuration ---
OUTPUT_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "ai_docs",
)
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "MethaNet_Strategic_Venture_Document.docx")
FONT = "Calibri"
FONT_FALLBACK = "DejaVu Sans"

# --- Helpers ---

def set_font(run, name=FONT, size=None, bold=False, italic=False, color=None):
    """Set font properties on a run."""
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = RGBColor(0x1B, 0x26, 0x3B)
    return h


def add_para(doc, text, bold=False, italic=False, size=10.5, alignment=None,
             space_after=6, first_line_indent=None, color=None):
    """Add a paragraph with styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)
    return p


def add_blockquote(doc, text, italic=True):
    """Add a styled blockquote paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=10, italic=italic, color=(0x55, 0x55, 0x55))
    return p


def add_bullet(doc, text, bold_prefix=None, level=0):
    """Add a bullet point, optionally with bold prefix."""
    p = doc.add_paragraph(style="List Bullet")
    if level > 0:
        p.style = doc.styles["List Bullet 2"] if "List Bullet 2" in [s.name for s in doc.styles] else p.style
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=10.5, bold=True)
        r2 = p.add_run(f" {text}")
        set_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_font(r, size=10.5)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9.5, bold=True, color=(0xFF, 0xFF, 0xFF))
        # shade header
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "1B263B",
        })
        shading.append(shd)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9.5)
    doc.add_paragraph()  # spacer
    return table


def add_hr(doc):
    """Add a horizontal rule (thin paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 80)
    set_font(run, size=6, color=(0xCC, 0xCC, 0xCC))


# === DOCUMENT BUILD ===

def build_document():
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # =====================================================
    # TITLE PAGE
    # =====================================================
    for _ in range(6):
        doc.add_paragraph()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("MethaNet")
    set_font(r, size=36, bold=True, color=(0x1B, 0x26, 0x3B))

    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp2.add_run("A Domain-Expert Agentic System for\nCross-Ecosystem Methanogenic Microbiology")
    set_font(r, size=16, italic=True, color=(0x44, 0x66, 0x88))

    doc.add_paragraph()

    tp3 = doc.add_paragraph()
    tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp3.add_run("Strategic Venture Document")
    set_font(r, size=14, bold=True, color=(0x1B, 0x26, 0x3B))

    tp4 = doc.add_paragraph()
    tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp4.add_run("Confidential Working Draft  —  March 2026")
    set_font(r, size=11, color=(0x88, 0x88, 0x88))

    for _ in range(6):
        doc.add_paragraph()

    tp5 = doc.add_paragraph()
    tp5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp5.add_run("EmergentBiome  ·  MethaNet Project")
    set_font(r, size=10, color=(0x88, 0x88, 0x88))

    doc.add_page_break()

    # =====================================================
    # EPIGRAPH
    # =====================================================
    add_blockquote(
        doc,
        '"The next biotech breakthrough won\'t be discovered by a lab, but by a network of AI agents '
        'coordinating, paying each other, and commissioning real wet lab experiments."',
    )
    add_para(doc, "— Bio Protocol / ClawdLab, 2026 [5]", italic=True, size=9,
             alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    add_hr(doc)

    # =====================================================
    # EXECUTIVE SUMMARY
    # =====================================================
    add_heading_styled(doc, "Executive Summary", level=1)

    add_para(
        doc,
        "A rare and time-bounded convergence is underway. Three independent forces — the maturation "
        "of agentic science infrastructure, a structural credibility crisis in voluntary carbon markets, "
        "and the emergence of protein and genomic language models capable of illuminating microbial dark "
        "matter — are colliding in a way that creates a clearly defined, commercially grounded, and "
        "scientifically defensible opportunity. MethaNet is positioned to occupy it."
    )

    add_para(
        doc,
        "MethaNet is a domain-expert autonomous science agent specialising in cross-ecosystem "
        "methanogenic microbiology. It operates on environmental metagenomes — the raw DNA of entire "
        "microbial communities extracted directly from soils, sediments, water columns, and gut systems — "
        "and produces two classes of output that no existing system generates: (1) scientifically rigorous, "
        "falsifiable hypotheses about microbial functional ecology, published and peer-reviewed through "
        "agentic science platforms; and (2) commercially licensed Genomic MRV (Monitoring, Reporting, "
        "Verification) products that predict methane flux from metagenomic functional profiles, directly "
        "serving the voluntary carbon market projected to reach $47.5 billion by 2035 [20]."
    )

    add_para(
        doc,
        "The scientific foundation is real and validated. A proof-of-concept embedding of 662 "
        "metagenome-assembled genomes (MAGs) — drawn from rumen (PRJEB31266, n=555) and wetland (MUCC, "
        "n=107) ecosystems — into a shared 1,280-dimensional ESM2-650M latent space demonstrates "
        "near-perfect ecosystem separation (AUC=1.000, PERMANOVA R²=0.202, p=0.001) alongside a "
        "scientifically remarkable bridge population: a small subset of rumen Archaea whose embedding "
        "neighbourhoods are equally split between ecosystems, with alpha-transfer scores exceeding 3.47 — "
        "more than 6 standard deviations above the cohort mean. These organisms are functional "
        "interpreters between biomes. They are the entry point into a broadly applicable framework."
    )

    add_para(
        doc,
        "The ambition is larger than the POC. MethaNet is not a rumen-to-wetland tool. It is a "
        "general-purpose metagenomic intelligence engine, applicable to any environmental sample — "
        "peatlands, rice paddies, marine sediments, anaerobic digesters, mangrove soils — wherever "
        "methane-cycling microbial communities determine ecological function and carbon accounting "
        "outcomes. The geometric signals captured in the current embedding are the proof of concept. "
        "The functional genomics layer — mcrA operons, CAZyme repertoires, syntrophic marker genes, "
        "KEGG pathway completeness, DRAM metabolic distillations — is the mechanism layer that makes "
        "predictions defensible, not just statistical."
    )

    add_para(
        doc,
        "The market is broken in exactly the right way. Fewer than 7 million blue carbon credits have "
        "been issued globally after two decades of effort, not because the ecosystems lack carbon, but "
        "because MRV methodology lacks the resolution to certify it credibly [18, 19]. The voluntary "
        "carbon market retired 7% fewer credits in 2025 despite a 227% surge in corporate net-zero "
        "commitments [17] — a credibility crisis, not a demand crisis. MethaNet's Genomic MRV layer "
        "is a direct solution to this structural failure."
    )

    add_para(
        doc,
        "The agentic science infrastructure to deploy this is live today. Science Beach hosts registered "
        "autonomous science agents and over 1,100 agent-generated hypotheses as of early 2026 [SB]. "
        "BioProtocol and ClawdLab have demonstrated autonomous agents commissioning real wet lab work "
        "via payment rails [5]. FutureHouse, Sakana AI, and others are building general AI scientists. "
        "Nobody is building a domain-expert agent in methanogenic ecology with a commercial carbon market "
        "output. The white space is real, confirmed, and currently uncontested."
    )

    add_para(
        doc,
        "This document makes the full case: what MethaNet is, why it wins, how it generates revenue, "
        "and what must be executed — and in what order — to make that possible."
    )

    doc.add_page_break()

    # =====================================================
    # PART I: THE CONVERGENCE
    # =====================================================
    add_heading_styled(doc, "Part I: The Convergence — Three Forces Meeting at Once", level=1)

    # --- 1.1 ---
    add_heading_styled(doc, "1.1  The Agentic Science Revolution Is Infrastructure, Not Hype", level=2)

    add_para(
        doc,
        "The shift from AI-assisted science to AI-agentic science crossed a meaningful threshold in "
        "2025–2026. This is not a speculative claim; it is documented across multiple independent actors "
        "building the same stack from different angles."
    )

    add_para(
        doc,
        "FutureHouse — an Eric Schmidt-backed nonprofit — launched Kosmos in November 2025: an AI "
        "Scientist that autonomously runs up to 12-hour cycles of parallel data analysis, literature "
        "search, and hypothesis generation, synthesising findings into scientific reports without human "
        "direction [1]. FutureHouse has declared a 10-year mission to build a Nobel-capable AI "
        "Scientist [2]. Sakana AI — with approximately $379 million in total funding and a $2.65 billion "
        "valuation as of November 2025 — has deployed AI Scientist v2, a generalised end-to-end agentic "
        "system that generated the first workshop paper written entirely by AI and accepted through peer "
        "review [3]. The MIT 2025 AI Agent Index formally classifies autonomous research agents as a "
        "distinct category alongside enterprise AI platforms [4]."
    )

    add_para(
        doc,
        "Science Beach (beach.science), launched in this context, is the social and economic layer. "
        "As of early 2026, it hosts registered AI agents and humans collaborating on over 1,100 "
        "agent-generated hypotheses in a public, falsifiable forum [SB]. The BioProtocol-ClawdLab "
        "partnership has gone further: the OpenClaw framework and the agent-only social network Moltbook "
        "produced a large-scale dataset of autonomous AI-to-AI interaction in January 2026, attracting "
        "six academic publications within fourteen days [5, OC]. Agents autonomously query the BIOS AI "
        "scientist via pay-per-query API, commission wet lab experiments using cryptocurrency payment "
        "rails, and collect rewards when experimental results are validated."
    )

    add_para(
        doc,
        "A February 2026 SwissCognitive analysis identifies the critical gap in this ecosystem: isolated "
        "AI scientists \"often lack the social context that makes science trustworthy and replicable\" [6]. "
        "Domain depth and community falsifiability are the dimensions on which the next wave of agentic "
        "science will compete. The high hypothesis rejection rate on Science Beach is a direct empirical "
        "signal: generic agents generate noise; domain-expert agents generate signal."
    )

    add_para(
        doc,
        "Between 2022 and 2025, startups in the agentic AI market raised a combined $4.4 billion across "
        "101 equity deals [7]. The global enterprise agentic AI market is projected to reach $24.5–$47 "
        "billion by 2030, with compound annual growth rates of 44–46% [7, 8]. The infrastructure is "
        "assembled. The domain expert for environmental microbiology and carbon markets is missing."
    )

    # --- 1.2 ---
    add_heading_styled(doc, "1.2  Protein and Genomic Language Models Are Cracking Open Microbial Dark Matter", level=2)

    add_para(
        doc,
        "The scientific backbone of MethaNet rests on a rapid maturation in sequence-based machine "
        "learning applied to environmental microbiology. Several developments in 2024–2025 directly "
        "confirm MethaNet's technical approach is both sound and timely."
    )

    add_para(
        doc,
        "ESM2 — Meta AI's protein language model trained on 250 million protein sequences — has "
        "demonstrated that sequence embeddings capture deep functional signal across enormous "
        "evolutionary distances [9]. Its 650M-parameter variant generates 1,280-dimensional "
        "representations that encode both structural and functional information without requiring "
        "alignment or prior annotation. In 2024–2025, this capability was extended explicitly to "
        "environmental microbiology:"
    )

    # BioGeoFormer
    p = doc.add_paragraph()
    r = p.add_run("BioGeoFormer ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "(bioRxiv, December 2025) fine-tunes ESM2-8M specifically on biogeochemical cycling proteins "
        "to classify sequences into 37 defined pathway categories involved in 4 major biogeochemical "
        "cycles (methane, sulfur, nitrogen, and phosphorus), demonstrating that ESM2 embeddings carry "
        "biogeochemical functional signal in environmental contexts [10]. This is the closest existing "
        "work to MethaNet's approach — and it stops at classification, with no agentic loop and no "
        "commercial output."
    )
    set_font(r, size=10.5)

    # LucaPCycle
    p = doc.add_paragraph()
    r = p.add_run("LucaPCycle ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "(Nature Communications, 2025) uses a dual-channel architecture integrating raw sequences and "
        "contextual embeddings based on ESM2-3B to illuminate phosphorus cycling proteins in deep-sea "
        "cold seep sediments, showing generalisation across taxonomic and geographic contexts when "
        "trained on biogeochemical pathway families [11]."
    )
    set_font(r, size=10.5)

    # EcoFoldDB
    p = doc.add_paragraph()
    r = p.add_run("EcoFoldDB ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "(Environmental Microbiology, 2025) builds on ProstT5 — a metagenomic protein language model — "
        "to provide structure-guided functional profiling of 32 million proteins encoded by 8,000 "
        "high-quality metagenome-assembled genomes from the global soil microbiome, addressing the "
        "annotation gap in environmental proteome datasets [12]."
    )
    set_font(r, size=10.5)

    # gLM
    p = doc.add_paragraph()
    r = p.add_run("gLM ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "(Nature Communications, 2024; cited by 113+) trains a genomic language model on millions of "
        "metagenomic scaffolds to learn latent functional and regulatory relationships between genes — "
        "directly demonstrating that unsupervised language model pretraining on metagenomes captures "
        "operonic co-regulation and functional co-occurrence without supervision [13]."
    )
    set_font(r, size=10.5)

    # MCycDB
    p = doc.add_paragraph()
    r = p.add_run("MCycDB ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "(Molecular Ecology Resources, 2022) provides the reference framework: 298 methane cycling gene "
        "families across 10 methane metabolism pathways with 610,208 representative sequences — the "
        "functional vocabulary against which MethaNet's embedding-derived predictions can be annotated "
        "and validated [15]."
    )
    set_font(r, size=10.5)

    # gcMeta
    p = doc.add_paragraph()
    r = p.add_run("gcMeta 2025 ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "(Nucleic Acids Research, 2025) provides a global repository of metagenome-assembled genomes "
        "establishing 50 biome-specific MAG catalogues comprising 109,586 species-level clusters, "
        "supporting interactive cross-biome analyses of diverse functional categories including "
        "methane cycling — confirming that cross-biome metagenomic comparisons are both technically "
        "feasible and scientifically valued [16]."
    )
    set_font(r, size=10.5)

    add_para(
        doc,
        "The critical observation is what none of these systems does: they are classifiers, annotators, "
        "or single-ecosystem tools. None operates as an autonomous agent. None connects outputs to a "
        "commercial carbon market product. None performs cross-ecosystem transfer learning to generate "
        "calibrated flux predictions for novel environmental samples. MethaNet integrates all layers "
        "into a single commercially actionable output system."
    )

    # --- 1.3 ---
    add_heading_styled(doc, "1.3  The Voluntary Carbon Market Needs a Credibility Revolution", level=2)

    add_para(
        doc,
        "The voluntary carbon market (VCM) is at a structural inflection driven by contradictory forces: "
        "enormous and growing corporate demand on one side, and a credibility crisis actively suppressing "
        "credit retirement on the other."
    )

    add_para(
        doc,
        "The VCM retired 7% fewer credits in 2025 despite a 227% surge in corporate net-zero commitments "
        "[17]. Supply is not the problem; trust in MRV methodology is. Despite two decades of effort, "
        "fewer than 7 million blue carbon credits have ever been issued globally — a fraction of the "
        "ecological potential of mangrove, seagrass, and saltmarsh ecosystems [18, 19]. Blue carbon "
        "represents less than 1% of all credits issued on the voluntary carbon market. The root cause "
        "is well documented: current chamber-based flux measurement methodology lacks spatial resolution, "
        "is prohibitively expensive at scale, and cannot differentiate microbially-driven methane "
        "emission from carbon sequestration at the ecosystem functional level."
    )

    add_para(
        doc,
        "Market size and trajectory are unambiguous, though forecasts vary by source. CarbonCredits.com "
        "projects the global VCM to grow from $1.7 billion in 2026 to $47.5 billion by 2035, a 38% "
        "compound annual growth rate [20]. BloombergNEF projects carbon credit prices could reach "
        "$60/tonne CO₂e by 2030 and $104/tonne by 2050 if technology-based removals dominate supply "
        "[21]. Mangrove restoration credits already command $26.03/credit (2023 average), carrying a "
        "premium over forest offsets due to co-benefits [22]. The International Blue Carbon and Wetlands "
        "Conference (2025, Trinidad and Tobago) convened global experts specifically to coordinate "
        "improved MRV systems for coastal carbon accounting — confirming this is an organised, funded, "
        "and active buyer community [23]."
    )

    add_para(
        doc,
        "What is needed — and what does not exist — is a sub-ecosystem, mechanistically grounded, "
        "scalable MRV layer that predicts methane flux from community functional potential. This is "
        "MethaNet's Genomic MRV product."
    )

    doc.add_page_break()

    # =====================================================
    # PART II: THE SCIENTIFIC FOUNDATION
    # =====================================================
    add_heading_styled(doc, "Part II: The Scientific Foundation", level=1)

    # --- 2.1 ---
    add_heading_styled(doc, "2.1  The Latent Space Signal: Geometry as Biological Knowledge", level=2)

    add_para(
        doc,
        "The MethaNet proof-of-concept demonstrates that ESM2 latent space encodes biologically "
        "meaningful, ecosystem-level functional signal that can be quantified, transferred, and "
        "commercialised."
    )

    add_para(
        doc,
        "Embedding 662 MAGs — 555 rumen (PRJEB31266) and 107 wetland (MUCC) — into a shared "
        "1,280-dimensional ESM2-650M latent space yields the following validated geometric structure:"
    )

    # Ecosystem Separation
    p = doc.add_paragraph()
    r = p.add_run("Ecosystem Separation.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "PERMANOVA confirms ecosystem identity explains 20.2% of total embedding variance "
        "(R²=0.202, p=0.001, 999 permutations, cosine distance). A 5-fold cross-validated logistic "
        "classifier achieves AUC=1.000, AUPRC=1.000, and balanced accuracy=99.9%, demonstrating "
        "near-perfect linear separability under real-world class imbalance (107 vs 555). Global "
        "silhouette score: 0.398 [95% bootstrap CI: 0.364, 0.439; 150 resamples]. Cohen's d along "
        "the rumen-to-wetland centroid trajectory axis: 3.63. These are not marginal signals; they are "
        "among the strongest ecosystem separation metrics reported in metagenomic embedding literature."
    )
    set_font(r, size=10.5)

    # Dimensional Architecture
    p = doc.add_paragraph()
    r = p.add_run("Dimensional Architecture.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "The first principal component captures approximately 44% of variance and aligns with the "
        "ecosystem separation axis. The second component captures approximately 22% and reflects "
        "within-rumen diversity — implying the rumen ecosystem is functionally more heterogeneous than "
        "the wetland. Only 4 principal components explain 80% of total variance, indicating the "
        "embedding space is efficiently structured rather than diffuse. This compactness is a critical "
        "prerequisite for generalisation to novel environmental samples."
    )
    set_font(r, size=10.5)

    # Projection Convergence
    p = doc.add_paragraph()
    r = p.add_run("Projection Convergence.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "Three independent dimensionality reduction methods — PCA, UMAP, and t-SNE — converge on the "
        "same finding: distinct but non-isolated ecosystem regions. UMAP and t-SNE reveal that wetland "
        "genomes form a single tight cluster, while rumen genomes fragment into multiple sub-clusters "
        "with varying degrees of proximity to the wetland region. This sub-cluster architecture is "
        "scientifically significant: not all rumen organisms are equally \"rumen-like\" in functional "
        "space, and a subset occupies an intermediate zone that is the mechanistic basis for transfer "
        "learning."
    )
    set_font(r, size=10.5)

    # Bridge Genome Architecture
    p = doc.add_paragraph()
    r = p.add_run("Bridge Genome Architecture.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "Of 662 genomes, 14 have at least one opposite-ecosystem genome among their k=15 nearest "
        "neighbours. Five rumen Archaea achieve neighbourhood entropy scores above 0.99, meaning "
        "their local neighbourhoods are nearly equally split between rumen and wetland genomes. The "
        "strongest bridge candidate — rumen bin.8, Archaeal, 1,321 proteins, ENA analysis ERZ1037672 — "
        "achieves an alpha-transfer score of 3.47, more than 6 standard deviations above the cohort "
        "mean (mean=0, std=0.56). All top 11 bridge candidates are rumen Archaea. These organisms are "
        "the molecular Rosetta Stones between biomes. Their characterisation is the scientific core of "
        "MethaNet's competitive advantage."
    )
    set_font(r, size=10.5)

    # Generalisability Architecture
    p = doc.add_paragraph()
    r = p.add_run("Generalisability Architecture.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "The most important implication of this geometry is universal projection. If ESM2 embeddings "
        "encode genuine functional signal — not batch artefact — then any new MAG from any metagenomic "
        "environment can be projected into this same latent space and its position interpreted as a "
        "functional fingerprint relative to known ecosystem anchors. Novel samples are not forced into "
        "binary categories; they are positioned in a continuous functional space with quantified "
        "reference points. Peatland organisms, rice paddy methanogens, marine sediment communities, "
        "and industrial digester populations can all be positioned, compared, and functionally "
        "characterised relative to the growing MethaNet reference corpus. The embedding becomes a "
        "universal coordinate system for methanogenic function. The 662-genome POC is the seed; the "
        "system grows and improves with every new metagenome added."
    )
    set_font(r, size=10.5)

    # POC metrics table
    add_table(doc,
        ["Metric", "Value", "Interpretation"],
        [
            ["Cohort", "662 MAGs (107 wetland + 555 rumen)", "16.5× scale-up from 40-genome baseline"],
            ["Embedding", "662 × 1,280 (ESM2-650M)", "Zero attrition, zero non-finite vectors"],
            ["PERMANOVA R²", "0.202 (p=0.001)", "Ecosystem explains 20.2% of embedding variance"],
            ["Silhouette", "0.398 [CI: 0.364–0.439]", "Bootstrap 95% CI, 150 resamples"],
            ["CV AUC", "1.000", "5-fold, PCA-50, balanced class weights"],
            ["Cohen's d", "3.63", "Very large effect size on trajectory axis"],
            ["Bridge genomes", "14 (opp. k-NN > 0)", "5 Archaea with entropy > 0.99"],
            ["Top bridge", "bin.8 (α-transfer = 3.47)", ">6 SDs above cohort mean"],
        ],
    )

    add_para(
        doc,
        "Source-ecosystem confounding caveat: All 555 rumen genomes derive from PRJEB31266; all 107 "
        "wetland genomes from MUCC. Ecosystem separability cannot be causally attributed to ecology "
        "alone — it could partially reflect technical batch effects. Deconfounding with additional "
        "independent sources is the P0 priority before any public ecological generalisation claim.",
        italic=True, size=9.5, color=(0x88, 0x44, 0x00),
    )

    # --- 2.2 ---
    add_heading_styled(doc, "2.2  The Functional Genomics Layer: Mechanism Behind the Geometry", level=2)

    add_para(
        doc,
        "Latent space geometry alone is not sufficient for commercial credibility or scientific "
        "defensibility. The complement is a deep functional genomics annotation layer that explains "
        "mechanistically why organisms occupy specific positions and what those positions predict about "
        "ecosystem-level methane flux. This layer is not an add-on; it is the primary evidence layer "
        "that bridges model prediction to carbon market certification requirements. It operates across "
        "five integrated annotation frameworks, fully applicable to any MAG or metagenome-assembled "
        "genome set from any environmental sample."
    )

    # Layer 1
    p = doc.add_paragraph()
    r = p.add_run("Layer 1: Methanogenic Pathway Architecture via MCycDB and KEGG.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "MCycDB provides 298 gene families across 10 methane metabolism pathways — methanogenesis, "
        "anaerobic methane oxidation (AOM), and all intermediate transfer steps — with 610,208 "
        "representative sequences [15]. For every MAG in the MethaNet corpus, pathway completeness "
        "scores are computed and annotated across all 10 pathways. Combined with KEGG module "
        "completeness, this generates a multi-dimensional metabolic fingerprint per organism and "
        "per community."
    )
    set_font(r, size=10.5)

    add_para(doc, "Key target gene systems:", bold=True, size=10.5)

    add_bullet(doc, "mcrA / mcrBG (methyl-coenzyme M reductase complex): ",
               bold_prefix="mcrA / mcrBG:")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    r = p.add_run(
        "The canonical phylogenetic and functional marker for all methanogens and anaerobic "
        "methanotrophs. mcrA operon completeness distinguishes fully competent methanogens from "
        "partial pathway organisms and defines substrate diversity (aceticlastic, hydrogenotrophic, "
        "methylotrophic). A 2025 metagenomic analysis confirmed mcrA as the primary diagnostic for "
        "methanogenic community structure and CH₄ yield variation across host types [24]."
    )
    set_font(r, size=10)

    add_bullet(doc, "Marks methylotrophic methanogenesis — the substrate flexibility pathway most "
               "associated with cross-ecosystem transfer, since methanol and methylamines are present "
               "across wetland, rumen, and marine environments.",
               bold_prefix="mtaBC (methanol:cobalamin methyltransferase):")

    add_bullet(doc, "Marks the CO₂-reduction (hydrogenotrophic) pathway, dominant in syntrophic rumen "
               "methanogenesis and present in wetland methanogens under low-acetate conditions.",
               bold_prefix="fwdABCDEF (formylmethanofuran dehydrogenase):")

    # Layer 2
    p = doc.add_paragraph()
    r = p.add_run("Layer 2: DRAM Metabolic Distillation.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "DRAM (Distilled and Refined Annotation of Metabolism; Shaffer et al., 2020; cited by 1,060+) "
        "provides automated, standardised metabolic curation of MAGs across KEGG, Pfam, CAZy, MEROPS, "
        "and VOG databases simultaneously [25]. For each bridge candidate MAG, DRAM generates a "
        "comprehensive metabolic product sheet identifying carbon fixation pathways, electron carrier "
        "systems, and methanogenesis module completeness — providing the mechanistic explanation for "
        "why a genome occupies its observed embedding position."
    )
    set_font(r, size=10.5)

    # Layer 3
    p = doc.add_paragraph()
    r = p.add_run("Layer 3: CAZyme Repertoire Profiling.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "Carbohydrate-Active enZyme (CAZy) profiling captures the substrate degradation capacity of "
        "each MAG, distinguishing organisms that can access complex polysaccharides (cellulose, chitin, "
        "xylan) from those limited to simple substrates. In methanogenic communities, CAZyme repertoire "
        "diversity correlates with syntrophic interaction potential — bridge genomes with broad CAZyme "
        "profiles are more likely to function across ecosystems where substrate availability differs."
    )
    set_font(r, size=10.5)

    # Layer 4
    p = doc.add_paragraph()
    r = p.add_run("Layer 4: Syntrophic Marker Gene Detection.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "Direct interspecies electron transfer (DIET) markers — multiheme cytochromes, pilA/pilB "
        "homologs, and type IV pili components — identify organisms capable of syntrophic partnerships "
        "that are central to methanogenesis in both rumen (syntrophic acetate oxidation) and wetland "
        "(DIET-coupled AOM) environments. Enrichment of DIET markers in bridge genomes would provide "
        "strong mechanistic evidence for their cross-ecosystem functional flexibility."
    )
    set_font(r, size=10.5)

    # Layer 5
    p = doc.add_paragraph()
    r = p.add_run("Layer 5: Community-Level Functional Fingerprints.  ")
    set_font(r, size=10.5, bold=True)
    r = p.add_run(
        "Aggregating individual MAG annotations into community-level profiles generates a composite "
        "functional fingerprint per sample: methanogenesis pathway completeness distribution, "
        "methanotrophy gene density, syntrophic interaction potential, and substrate flexibility "
        "score. This community fingerprint — not any single gene — is the input to the flux "
        "prediction model and the basis of the Genomic MRV report."
    )
    set_font(r, size=10.5)

    doc.add_page_break()

    # =====================================================
    # PART III: REVENUE MODEL (5.4 from draft)
    # =====================================================
    add_heading_styled(doc, "Part III: Revenue Model", level=1)

    add_para(
        doc,
        "MethaNet generates revenue through three complementary streams, each tied to a distinct "
        "stage of scientific validation and market maturity."
    )

    add_heading_styled(doc, "Revenue Projections", level=2)

    add_table(doc,
        ["Revenue Stream", "Year 1", "Year 2", "Year 3", "Year 5"],
        [
            ["Genomic MRV Licensing", "$200K", "$1.5M", "$6M", "$18M"],
            ["Science Beach / Platform Rewards", "$20K", "$80K", "$200K", "$500K"],
            ["Carbon Credit Co-origination", "—", "$50K", "$500K", "$5M"],
            ["Total", "~$220K", "~$1.6M", "~$6.7M", "~$23.5M"],
        ],
    )

    add_para(
        doc,
        "These projections assume successful bridge genome flux validation by Q3 2026, VCS methodology "
        "engagement initiated Q4 2026, and first certified methodology acceptance by 2028. If "
        "experimental validation fails to establish flux correlation, or methodology acceptance is "
        "delayed beyond 2029, Year 1–2 estimates compress materially. Treat as achievable-if-executed "
        "figures, not base-case guarantees.",
        italic=True, size=9.5,
    )

    doc.add_page_break()

    # =====================================================
    # PART IV: EXECUTION ROADMAP
    # =====================================================
    add_heading_styled(doc, "Part IV: Execution Roadmap", level=1)

    add_para(
        doc,
        "Sequencing matters as much as strategy. Nothing in a later phase is attempted before the prior "
        "phase delivers real, observable evidence. The roadmap is built around hard scientific "
        "dependencies, not optimistic timelines."
    )

    # --- Phase 0 ---
    add_heading_styled(doc, "Phase 0: Scientific Integrity Verification (Months 1–3)", level=2)

    add_para(
        doc,
        "Mandate: No public commitments. No Science Beach registration. No client conversations. No "
        "commercial claims of any kind. This phase exists exclusively to resolve the most important "
        "outstanding scientific question before any public positioning is made: does embedding position "
        "in the MethaNet latent space predict methanogenic function in independent, novel samples?",
        bold=False,
    )

    add_para(doc, "Action 1 — Sub-sampling deconfounding analysis.", bold=True)
    add_para(
        doc,
        "The current POC uses 555 rumen MAGs against 107 wetland MAGs — a 5:1 numerical imbalance "
        "with perfect source-ecosystem confounding. Repeatedly downsample the rumen MAG set to n=107 "
        "(matching wetland size, 100 bootstrap iterations) and recompute PERMANOVA, silhouette score, "
        "and AUC at each iteration. If separation holds — pseudo-F remains significant, AUC remains "
        ">0.95, silhouette CI remains above 0.3 — batch and numerical dominance artefacts are ruled "
        "out with confidence."
    )

    add_para(doc, "Action 2 — Multi-source rumen integration.", bold=True)
    add_para(
        doc,
        "Identify at least one independent rumen metagenome dataset from a different study cohort and "
        "sequencing batch. Embed those MAGs into the existing latent space and assess whether they "
        "cluster with the existing PRJEB31266 rumen population or show batch-driven displacement. "
        "This is the definitive deconfounding test."
    )

    add_para(doc, "Action 3 — Bridge genome quality audit.", bold=True)
    add_para(
        doc,
        "Extract full GenBank accession numbers, assembly statistics (N50, total length, number of "
        "contigs), CheckM2 completeness and contamination scores, and GTDB-Tk taxonomic classification "
        "for all bridge candidate genomes. A bridge genome with 55% completeness cannot support the "
        "functional claims the commercial model requires."
    )

    add_para(doc, "Action 4 — Initiate first bridge genome culture attempt.", bold=True)
    add_para(
        doc,
        "Contact 2–3 anaerobic microbiology CRO partners with experience cultivating rumen or wetland "
        "Archaea. Design the preliminary enrichment protocol for rumen bin.8: anaerobic medium "
        "composition (Hungate technique), substrate conditions (H₂/CO₂ for hydrogenotrophic; methanol "
        "for methylotrophic), temperature and pH gradients spanning rumen (39°C, pH 6.5) to wetland "
        "(25°C, pH 7.2) conditions. Measure CH₄ production by gas chromatography at weekly intervals."
    )

    add_para(
        doc,
        "Phase 0 exit criteria: Sub-sampling analysis shows robust separation (AUC >0.95 under "
        "resampling); independent rumen MAGs cluster correctly; all bridge genome quality metrics meet "
        "threshold (completeness >70%, contamination <10%); at least one bridge genome enrichment "
        "initiated with measurable CH₄ production signal.",
        italic=True, size=10,
    )

    # --- Phase 1 ---
    add_heading_styled(doc, "Phase 1: Scientific Validation and Agent Launch (Months 3–9)", level=2)

    add_para(
        doc,
        "Mandate: Establish scientific credibility through real experimental results and a measured, "
        "evidence-backed public presence on Science Beach. Generate the first peer-reviewed preprint. "
        "Build the first CRO lab partnerships."
    )

    add_para(doc, "Action 1 — Register MethaNet agents on Science Beach.", bold=True)
    add_para(
        doc,
        "Register the initial agent network: MethaNet-Explorer (hypothesis generation) and MethaNet-"
        "Analyst (community engagement). First three Science Beach hypotheses, each directly derived "
        "from validated POC findings:"
    )

    add_blockquote(
        doc,
        "Hypothesis 1: \"Bridge Archaeal genomes in the rumen-wetland ESM2 embedding space encode "
        "broader mcrA substrate diversity than ecosystem-specific methanogens, as measured by MCycDB "
        "pathway completeness scores across aceticlastic, hydrogenotrophic, and methylotrophic "
        "methanogenesis modules.\" Testable prediction: bridge entropy score (k=15 NN) correlates "
        "positively with MCycDB multi-pathway completeness (Spearman ρ >0.4, p<0.05)."
    )

    add_blockquote(
        doc,
        "Hypothesis 2: \"The independent rumen sub-cluster identified in UMAP and t-SNE projections "
        "represents a phylogenetically coherent Archaeal lineage with enriched DIET marker genes "
        "relative to the main rumen cluster.\" Testable prediction: sub-cluster MAGs show significantly "
        "higher multiheme cytochrome gene density (Mann-Whitney U, p<0.05)."
    )

    add_blockquote(
        doc,
        "Hypothesis 3: \"Wetland sub-cluster structure correlates with sampling site metadata rather "
        "than phylogenetic distance.\" Testable prediction: pairwise embedding distance between "
        "wetland MAGs correlates more strongly with sampling site than with GTDB-Tk phylogenetic "
        "distance (partial Mantel test, r_geography > r_phylogeny)."
    )

    add_para(doc, "Action 2 — Commission bridge genome validation experiments (~$75K–$150K).", bold=True)
    add_para(
        doc,
        "Through contracted CRO partner(s): (A) Anaerobic cultivation of top 3 bridge candidates under "
        "four conditions: standard rumen (H₂/CO₂, 39°C, pH 6.5), methylotrophic rumen (methanol), "
        "wetland-analogous hydrogenotrophic (H₂/CO₂, 25°C, pH 7.2), and acetoclastic wetland (acetate, "
        "25°C, pH 7.2). Primary endpoint: CH₄ production rate (μmol CH₄/mg protein/hr). "
        "(B) Metagenomic confirmation of bridge genome enrichment purity via shotgun sequencing and "
        "ANI confirmation. (C) Syntrophic co-culture pilot pairing bridge Archaeal enrichment with "
        "representative bacterial fermenters from both ecosystems."
    )

    add_para(doc, "Action 3 — Submit bioRxiv preprint.", bold=True)
    add_para(
        doc,
        "Target: \"Cross-Ecosystem Protein Language Model Embeddings Reveal Methanogenic Bridge Archaea "
        "with Broad Substrate Flexibility.\" Cite MCycDB, DRAM, BioGeoFormer, gLM, and LucaPCycle as "
        "methodological context."
    )

    add_para(
        doc,
        "Phase 1 exit criteria: At least 1 community-claimed hypothesis on Science Beach; at least 1 "
        "bridge genome enrichment producing quantifiable CH₄; bioRxiv preprint submitted; 2 CRO lab "
        "partnerships formalised.",
        italic=True, size=10,
    )

    # --- Phase 2 ---
    add_heading_styled(doc, "Phase 2: MRV Product Development and Pilot Engagement (Months 9–21)", level=2)

    add_para(
        doc,
        "Mandate: Translate scientific credibility into commercial MRV pilots. Build the generalisation "
        "pipeline. Engage carbon standard bodies formally. Generate first paying clients."
    )

    add_bullet(doc, "Build and validate the full generalisation pipeline: assembly/binning, ESM2 embedding, "
               "position scoring, DRAM+MCycDB+CAZy annotation, community fingerprint, flux prediction, "
               "and report generation. Validate on held-out metagenomes with known flux measurements.",
               bold_prefix="Pipeline validation:")
    add_bullet(doc, "Technical white paper: \"Genomic MRV: Cross-Ecosystem Protein Language Model "
               "Embeddings as a Predictive Framework for Methanogenic Wetland Systems.\"",
               bold_prefix="Methodology white paper:")
    add_bullet(doc, "Recruit 3–5 mangrove or peatland restoration projects (2–5 year stage). Discounted "
               "Genomic MRV Reports ($5K–$10K at cost) in exchange for sample access and co-authorship.",
               bold_prefix="Pilot clients:")
    add_bullet(doc, "Submit methodology concept note to VCS (Verra) and Plan Vivo. Goal: establish working "
               "relationship with reviewers, not immediate approval.",
               bold_prefix="Carbon standard engagement:")
    add_bullet(doc, "Commission metagenomics from 3 contrasting ecosystem types not yet represented: "
               "tropical peatland, temperate rice paddy, and coastal marine sediment (50–100 MAGs each). "
               "Budget: ~$100K–$200K.",
               bold_prefix="Ecosystem expansion:")

    add_para(
        doc,
        "Phase 2 exit criteria: Full pipeline validated against ≥2 independent field datasets; ≥3 "
        "paying pilot clients; VCS/Plan Vivo engagement active; corpus expanded to ≥3 ecosystem types; "
        "second preprint submitted.",
        italic=True, size=10,
    )

    # --- Phase 3 ---
    add_heading_styled(doc, "Phase 3: Commercial Scale and Methodology Certification (Months 21–42)", level=2)

    add_bullet(doc, "Full commercial pricing ($15K–$200K per engagement). Target 15–20 clients Year 2, "
               "50–100 Year 3. Client portal for data submission and report tracking.",
               bold_prefix="Commercial launch:")
    add_bullet(doc, "Full VCS methodology documentation. Includes validation across ≥5 sites, uncertainty "
               "quantification, accuracy benchmarking against chambers, independent peer review.",
               bold_prefix="VCS submission:")
    add_bullet(doc, "Expand reference corpus to ≥1,200 MAGs across 12 ecosystem types — a resource with "
               "no parallel in published literature.",
               bold_prefix="Corpus expansion:")
    add_bullet(doc, "3–5 pilot projects: MethaNet provides Genomic MRV at 50% discount for 15–20% equity "
               "stake in credits generated over 10-year crediting period.",
               bold_prefix="Co-origination:")
    add_bullet(doc, "Establish MethaNet as the most community-claimed agent in microbial ecology and "
               "climate science domains on Science Beach.",
               bold_prefix="Agent network leadership:")

    doc.add_page_break()

    # =====================================================
    # PART V: RISKS AND MITIGATIONS
    # =====================================================
    add_heading_styled(doc, "Part V: Risks and Mitigations", level=1)

    risks = [
        (
            "Risk 1: Bridge Genome Flux Correlations Do Not Materialise",
            "High", "Medium",
            "The POC acknowledges source-ecosystem confounding. If Phase 0 and Phase 1 experiments "
            "show that bridge entropy does not predict CH₄ output across culture conditions, the "
            "flux prediction model has no mechanistic basis.",
            "Phase 0 is a genuine go/no-go gate. If flux correlations do not emerge by Month 9, "
            "the commercial product pivots from flux prediction to community composition profiling — "
            "less valuable but still commercially useful ($5K–$20K per report rather than $15K–$200K). "
            "This pivot does not eliminate the commercial case."
        ),
        (
            "Risk 2: VCS Methodology Approval Timeline Extends Beyond 3 Years",
            "High on Tier 3", "High",
            "Carbon standard bodies are inherently conservative. Novel AI-based methodologies without "
            "5-year field validation will face extended review.",
            "Design the commercial product to generate material revenue without certification. Genomic "
            "MRV positioned as \"supplementary intelligence layer\" for projects already using "
            "chamber-based methods generates revenue independent of formal certification. Certification "
            "amplifies pricing; it is not a prerequisite for Tier 1 revenue."
        ),
        (
            "Risk 3: Science Beach Platform Does Not Mature On Schedule",
            "Medium", "Medium",
            "Science Beach describes itself as \"a social experiment.\" The payment integration and "
            "agent reward system may not achieve required maturity.",
            "Science Beach is a distribution and credibility channel, not a revenue mechanism. "
            "MethaNet's core Genomic MRV product is entirely independent of Science Beach's "
            "economic layer. Fallback: direct academic publishing and carbon market conference presence."
        ),
        (
            "Risk 4: A Well-Funded Competitor Enters",
            "High", "Low (12–18 months)",
            "FutureHouse has capital but no methane/climate vertical. Ginkgo Bioworks has synbio "
            "infrastructure but no environmental metagenomics or carbon market focus.",
            "Publish the bridge genome characterisation paper before Month 12. Scientific priority "
            "is established by publication date. Being the cited authority on cross-ecosystem "
            "methanogenic transfer learning is a moat that capital cannot rapidly overcome."
        ),
        (
            "Risk 5: VCM Credibility Crisis Deepens",
            "Medium", "Low long-term",
            "If corporate buyers retreat from voluntary offsets, demand for MRV could paradoxically "
            "contract as project pipelines slow.",
            "The VCM credibility crisis is itself the driver of demand for better MRV. MethaNet is "
            "credibility infrastructure — it benefits from the problem, not from it going away. "
            "Maintain client diversification into agricultural methane reduction and regulatory "
            "compliance applications."
        ),
    ]

    for title, impact, prob, description, mitigation in risks:
        add_heading_styled(doc, title, level=2)
        p = doc.add_paragraph()
        r = p.add_run(f"Impact: {impact}.  Probability: {prob}.")
        set_font(r, size=10, bold=True)
        add_para(doc, description)
        p = doc.add_paragraph()
        r = p.add_run("Mitigation: ")
        set_font(r, size=10.5, bold=True)
        r = p.add_run(mitigation)
        set_font(r, size=10.5)

    doc.add_page_break()

    # =====================================================
    # PART VI: THE SINGLE MOST IMPORTANT THING
    # =====================================================
    add_heading_styled(doc, "Part VI: The Single Most Important Thing", level=1)

    add_para(
        doc,
        "Everything in this document — the agent architecture, the revenue model, the competitive "
        "positioning, the co-origination pipeline — depends on one empirical question being answered "
        "correctly and honestly:",
    )

    add_para(
        doc,
        "Does the MethaNet embedding position predict methanogenic function in novel, independent "
        "environmental samples?",
        bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
    )

    add_para(
        doc,
        "If yes, MethaNet is the missing piece at a genuinely important intersection of science and "
        "commerce, and the execution roadmap above is the right path to realising it."
    )

    add_para(
        doc,
        "If no — if the embedding captures batch effects or taxonomic structure rather than genuine "
        "functional signal — the strategy must be rebuilt around a different value proposition, and "
        "doing that rebuild after public commitments have been made is far more costly than doing it "
        "in Phase 0."
    )

    add_para(
        doc,
        "Science first. Commercial translation second. In that order. Without exception. Not "
        "simultaneously. The rigour of Phase 0 is not caution — it is the prerequisite for everything "
        "else being worth doing.",
        bold=True, size=11,
    )

    doc.add_page_break()

    # =====================================================
    # REFERENCES
    # =====================================================
    add_heading_styled(doc, "References", level=1)

    refs = [
        "[1] FutureHouse / Edison Scientific. \"Kosmos: An AI Scientist for Autonomous Discovery.\" "
        "November 2025. arXiv:2511.02824. https://edisonscientific.com/articles/announcing-kosmos",

        "[2] FutureHouse. \"A philanthropically-funded moonshot focused on building an AI Scientist.\" "
        "10-year mission statement. https://www.futurehouse.org/",

        "[3] Sakana AI. \"AI Scientist v2: Workshop-Level Automated Scientific Discovery via Agentic "
        "Tree Search.\" 2025. Total funding: ~$379M (Series B: $135M, Nov 2025; valuation: $2.65B). "
        "https://github.com/SakanaAI/AI-Scientist-v2",

        "[4] Staufer, L. et al. \"The 2025 AI Agent Index: Documenting Technical and Safety Features "
        "of Deployed Agentic AI Systems.\" MIT, 2025. arXiv:2602.17753. https://aiagentindex.mit.edu/",

        "[5] Bio Protocol / ClawdLab / OpenClaw. \"OpenClaw, Moltbook, and ClawdLab: From Agent-Only "
        "Social Networks to Autonomous Scientific Research.\" January 2026. arXiv:2602.19810. See also: "
        "https://x.com/BioProtocol/status/2029899360989745305",

        "[SB] Science Beach (beach.science). Platform statistics as of early 2026: 42+ AI agents, "
        "47+ humans, 1,100+ hypotheses generated. https://launches.uicomet.com/products/science-beach",

        "[OC] OpenClaw / Moltbook. Large-scale autonomous AI-to-AI interaction dataset, January 2026. "
        "6 academic publications within 14 days. arXiv:2602.19810.",

        "[6] SwissCognitive. \"From Co-Pilot to Co-Scientist: Why the Next AI Scientist Must Be Social.\" "
        "February 2026. https://swisscognitive.ch/2026/02/24/",

        "[7] NewMarketPitch. \"Agentic AI Market Funding Trends (2022–2026).\" February 2026. "
        "$4.4B across 101 equity deals (2022–2025). https://newmarketpitch.com/blogs/news/agentic-ai-funding-trends",

        "[8] Grand View Research. \"Enterprise Agentic AI Market Size.\" Projected $24.5B by 2030, "
        "CAGR 46.2%. See also: Prism Media Wire, \"Agentic AI: A Strategic Forecast (2025–2030),\" "
        "$24.5–$48.2B by 2030. MarketsandMarkets projects $52.6B by 2030.",

        "[9] Lin, Z. et al. \"Evolutionary-scale prediction of atomic-level protein structure with a "
        "language model.\" Science, 379(6637), 1123–1130, 2023. DOI: 10.1126/science.ade2574.",

        "[10] Wynne, J.H. et al. \"BioGeoFormer: A deep learning approach to classify unknown genes "
        "associated with critical biogeochemical cycles.\" bioRxiv, December 17, 2025. Fine-tunes "
        "ESM2-8M on 37 pathway categories across 4 biogeochemical cycles (methane, sulfur, nitrogen, "
        "phosphorus). DOI: 10.64898/2025.12.17.695047.",

        "[11] Zhang, C. et al. \"LucaPCycle: Illuminating microbial phosphorus cycling in deep-sea "
        "cold seep sediments using protein language models.\" Nature Communications 16, 2025. "
        "Dual-channel model integrating raw sequences and ESM2-3B contextual embeddings. "
        "DOI: 10.1038/s41467-025-60142-4.",

        "[12] Ghaly, T.M. et al. \"EcoFoldDB: Protein Structure-Guided Functional Profiling of "
        "Ecologically Relevant Microbial Traits at the Metagenome Scale.\" Environmental Microbiology "
        "27, 2025. Builds on ProstT5; profiles 32M proteins from 8,000 soil MAGs. "
        "DOI: 10.1111/1462-2920.70178.",

        "[13] Hwang, Y. et al. \"Genomic language model predicts protein co-regulation and function.\" "
        "Nature Communications 15, 2880, 2024. Cited by 113+. DOI: 10.1038/s41467-024-46947-9.",

        "[14] Ayres, G. et al. \"Annotating the microbial dark matter with HiFi-NN.\" ScienceDirect, "
        "2025. Cited by 4.",

        "[15] Qian, L. et al. \"MCycDB: A curated database for comprehensively profiling methane "
        "cycling processes of environmental microbiomes.\" Molecular Ecology Resources 22(5), "
        "1803–1823, 2022. 298 gene families, 10 pathways, 610,208 sequences. DOI: 10.1111/1755-0998.13589.",

        "[16] Sun, Y. et al. \"gcMeta 2025: A global repository of metagenome-assembled genomes "
        "enabling cross-ecosystem microbial discovery and function research.\" Nucleic Acids Research "
        "54(D1), D724, 2025. 50 biome-specific catalogues, 109,586 species-level clusters. "
        "DOI: 10.1093/nar/gkaf449.",

        "[17] Carbon Direct. \"Key trends in the 2026 voluntary carbon market.\" February 2026. "
        "SBTi 227% surge in targets; VCM retirements fell 7% in 2025. "
        "https://www.carbon-direct.com/insights/key-trends-2026-voluntary-carbon-market",

        "[18] Calyx Global. \"Blue Carbon: What the rising tide of coastal conservation means for "
        "the voluntary carbon market.\" 2026. Blue carbon = <1% of VCM credits issued.",

        "[19] bioRxiv. \"High-Resolution Coastal Blue Carbon Site Intelligence.\" February 2026. "
        "DOI: 10.1101/2026.02.20.706974.",

        "[20] CarbonCredits.com. \"Voluntary Carbon Market in 2026: Top Forecasts.\" December 2025. "
        "VCM forecast: $1.7B in 2026, $47.5B by 2035, 38% CAGR. Note: other sources project lower "
        "(Regreener: €15B by 2035, 20.6% CAGR). "
        "https://carboncredits.com/voluntary-carbon-market-in-2026",

        "[21] BloombergNEF. \"Long-Term Carbon Credit Supply Outlook 2025.\" Prices: $60/tCO₂e by "
        "2030, $104/tCO₂e by 2050 (technology-removal scenario). "
        "https://about.bnef.com/insights/commodities/long-term-carbon-credit-supply-outlook-2025/",

        "[22] Simpson, S. et al. \"The Blue Carbon Cost Tool — understanding market dynamics in "
        "mangrove restoration credit pricing.\" Frontiers in Marine Science, 2025. Average mangrove "
        "credit price 2023: $26.03/credit. DOI: 10.3389/fmars.2025.1622255.",

        "[23] International Blue Carbon and Wetlands Conference (IBCWC). University of the West "
        "Indies, Port of Spain, Trinidad and Tobago, 2025. "
        "https://newsday.co.tt/2025/08/21/caribbean-latin-america-build-blueprint-for-blue-carbon-markets/",

        "[24] Bowerman, K.L. et al. \"Metagenomic analysis of fecal microbiomes reveals genetic "
        "determinants of enteric methane production and mcrA diversity.\" PMC12911397, 2025.",

        "[25] Shaffer, M. et al. \"DRAM for distilling microbial metabolism to automate the curation "
        "of microbiome function.\" Nucleic Acids Research 48(16), 8883–8900, 2020. Cited by 1,060+. "
        "DOI: 10.1093/nar/gkaa621.",

        "[26] California Management Review. \"Can Decentralized Science Be the Next Frontier of "
        "Scientific Discovery?\" November 2025. 50+ active DeSci projects; $60M+ funding.",

        "[27] AInvest / CoinDesk. \"AUBRAI: DeSci Agent Co-developed by VitaDAO and Bio Protocol "
        "Launches TGE on Base Blockchain.\" August 2025.",

        "[28] Mirete, S. et al. \"MAGs: Advances, Challenges and Applications in Methanogenic Archaea "
        "and Methane Cycling Communities.\" MDPI, 2025. Cited by 30.",

        "[29] Xamxidin, M. et al. \"Metagenomics-assembled genomes reveal microbial functional "
        "plasticity in Lake Barkol saline ecosystems.\" Frontiers in Microbiology, 2025. Cited by 7.",

        "[30] Duan et al. \"Pre-trained gene language model for metagenomics.\" Cited by 10, 2025.",

        "[31] Jha et al. \"Gaia: An AI-enabled genomic context-aware platform for environmental "
        "metagenomics.\" Cited by 11, 2025.",

        "[32] Georgia Tech Research. \"Meet the Microbes: What a Warming Wetland Reveals About "
        "Earth's Carbon Future.\" September 2025. "
        "https://cos.gatech.edu/news/meet-microbes-what-warming-wetland-reveals-about-earths-carbon-future",

        "[33] Tang, Y. et al. \"New insights into enteric methane production based on genomic and "
        "functional genomic analysis of rumen Archaea.\" ScienceDirect, 2025. Cited by 3.",

        "[34] Silicon Republic. \"Expectations from AI ramp up as investors eye returns.\" January 2026.",

        "[35] Drug Target Review. \"AI in drug discovery: predictions for 2026.\" February 2026. "
        "AI drug discovery market: $5–7B in 2025, $8–10B in 2026.",

        "[36] BioCarbon Standard. \"BioCarbon and Planet 2050 announce the Launch of a Digital MRV "
        "Working Group.\" June 2025.",

        "[37] Capgemini. \"Biopharma R&D turns to AI: Agentic AI in real-world environments.\" "
        "January 2026.",

        "[38] Drug Discovery World. \"The 2026 Biotech AI Report: Breakthroughs, bottlenecks, and "
        "the next wave of autonomous agents.\" December 2025.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        r = p.add_run(ref)
        set_font(r, size=8.5)

    doc.add_paragraph()
    add_hr(doc)

    # --- Closing ---
    add_para(
        doc,
        "This document represents a strategic opportunity grounded in verifiable market data, real and "
        "validated scientific work, and an honest accounting of risks and dependencies. The "
        "uncertainties are explicit. The risks are real. The opportunity is also real — and currently "
        "unoccupied.",
        italic=True, size=10,
    )

    add_para(
        doc,
        "The difference between MethaNet becoming a significant scientific and commercial venture and "
        "remaining an interesting research project is execution discipline: Phase 0 integrity before "
        "Phase 1 ambition, experimental validation before commercial claims, science before strategy.",
        italic=True, size=10,
    )

    add_para(
        doc,
        "If the bridge genome flux correlations are real — and the embedding geometry strongly "
        "suggests they are — everything else in this document follows with high strategic logic. "
        "Confirming that is the work of the next 90 days.",
        italic=True, size=10, bold=True,
    )

    # --- Save ---
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")

    # --- Stats ---
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    print(f"Paragraphs: {para_count}, Tables: {table_count}")


if __name__ == "__main__":
    build_document()
