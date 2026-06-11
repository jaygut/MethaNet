"""Build the Blue Catalyst deep-dive report v2.0.

The v2.0 report keeps the validated 662-genome ESM2 POC narrative, but
adds a stricter scientific critique, denominator accounting, source-confounding
diagnostics, bridge triage, and a functional-metagenomics roadmap aligned with
MethaNet's transfer-learning and methane MRV goals.

Usage:
    python scripts/build_blue_catalyst_deep_dive_report_v2.py
"""

# ruff: noqa: I001
from __future__ import annotations

import json
import textwrap
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.gridspec as gridspec
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from scipy import stats


PROJECT = Path(__file__).resolve().parent.parent
SNAPSHOT = (
    PROJECT
    / "results"
    / "blue_catalyst_poc"
    / "interim_snapshots"
    / "apolo_full_20260228_080644_embed_20260305_061952_notebook_interim_"
    "20260306_055012"
)
ARTIFACTS = (
    PROJECT
    / "results"
    / "blue_catalyst_poc"
    / "runs"
    / "apolo_full_20260228_080644_embed_20260305_061952"
    / "artifacts"
)
ANALYTICS = SNAPSHOT / "analytics"
TABLES = SNAPSHOT / "tables"
OUT_DIR = SNAPSHOT / "deep_dive_report_v2"
FIG_DIR = OUT_DIR / "figures"
REPORT = OUT_DIR / "blue_catalyst_deep_dive_report_v2.docx"
REVIEW_MD = OUT_DIR / "blue_catalyst_deep_dive_report_v2_review.md"
SOURCE_MANIFEST = OUT_DIR / "source_manifest.tsv"

FIG_DIR.mkdir(parents=True, exist_ok=True)


COLORS = {
    "rumen": "#A63D40",
    "wetland": "#087E8B",
    "archa": "#273043",
    "unknown": "#F2C14E",
    "bacteria": "#6A8D73",
    "ink": "#1E2329",
    "muted": "#697386",
    "paper": "#FAFBFC",
    "panel": "#FFFFFF",
    "gold": "#F4A261",
    "blue": "#2A6F97",
    "green": "#2A9D8F",
    "red": "#B23A48",
    "violet": "#6D597A",
}


def load_inputs() -> dict[str, object]:
    """Load all authoritative report inputs."""
    with np.load(ARTIFACTS / "genome_embeddings.npz", allow_pickle=True) as payload:
        embeddings = payload["embeddings"].astype(np.float32)

    projection = pd.read_csv(TABLES / "embedding_projection_partial.tsv", sep="\t")
    bridge_top = pd.read_csv(ARTIFACTS / "bridging_genomes_top.tsv", sep="\t")
    metadata = pd.read_csv(ARTIFACTS / "embedding_metadata.tsv", sep="\t")
    source_counts = pd.read_csv(ARTIFACTS / "sample_source_counts.tsv", sep="\t")
    stats_payload = json.loads((ARTIFACTS / "embedding_stats.json").read_text())
    summary = json.loads((ANALYTICS / "analytics_summary.json").read_text())

    return {
        "embeddings": embeddings,
        "projection": projection,
        "bridge_top": bridge_top,
        "metadata": metadata,
        "source_counts": source_counts,
        "stats": stats_payload,
        "summary": summary,
    }


def setup_plotting() -> None:
    sns.set_theme(style="white", context="paper", font_scale=1.05)
    plt.rcParams.update(
        {
            "figure.dpi": 220,
            "savefig.dpi": 320,
            "font.family": "DejaVu Sans",
            "axes.titleweight": "bold",
            "axes.labelcolor": COLORS["ink"],
            "xtick.color": COLORS["ink"],
            "ytick.color": COLORS["ink"],
        }
    )


def savefig(fig: plt.Figure, name: str) -> Path:
    path = FIG_DIR / name
    fig.savefig(path, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def add_panel_label(ax: plt.Axes, label: str, title: str) -> None:
    ax.set_title(f"{label}. {title}", loc="left", fontsize=12, pad=10)


def draw_box(
    ax: plt.Axes,
    xy: tuple[float, float],
    width: float,
    height: float,
    title: str,
    body: str,
    color: str,
    title_size: int = 10,
    body_size: int = 11,
) -> None:
    title_text = title if "\n" in title else textwrap.fill(title, max(10, int(width * 34)))
    body_text = body if "\n" in body else textwrap.fill(body, max(12, int(width * 48)))
    box = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.02,rounding_size=0.035",
        facecolor=color,
        edgecolor=COLORS["ink"],
        linewidth=1.2,
        alpha=0.95,
    )
    ax.add_patch(box)
    ax.text(
        xy[0] + width / 2,
        xy[1] + height * 0.65,
        title_text,
        ha="center",
        va="center",
        fontsize=title_size,
        fontweight="bold",
        color=COLORS["ink"],
        wrap=True,
    )
    ax.text(
        xy[0] + width / 2,
        xy[1] + height * 0.30,
        body_text,
        ha="center",
        va="center",
        fontsize=body_size,
        color=COLORS["ink"],
        wrap=True,
    )


def arrow(ax: plt.Axes, start: tuple[float, float], end: tuple[float, float]) -> None:
    ax.add_patch(
        FancyArrowPatch(
            start,
            end,
            arrowstyle="-|>",
            mutation_scale=16,
            linewidth=1.4,
            color=COLORS["muted"],
        )
    )


def safe_short(sample: str, max_len: int = 30) -> str:
    if len(sample) <= max_len:
        return sample
    return sample.replace("rumen__", "r__").replace("mucc__", "m__")[:max_len] + "..."


def make_figure_1(data: dict[str, object]) -> Path:
    """Evidence ladder, denominator integrity, and source confounding."""
    metadata = data["metadata"]
    source_counts = data["source_counts"]
    summary = data["summary"]
    stats_payload = data["stats"]

    pre_final_total = int(source_counts["n_samples"].sum())
    final_total = len(metadata)
    excluded = int(stats_payload.get("excluded_coassembly", pre_final_total - final_total))

    fig = plt.figure(figsize=(15, 10), facecolor=COLORS["paper"])
    gs = gridspec.GridSpec(2, 2, hspace=0.32, wspace=0.25)

    ax = fig.add_subplot(gs[0, 0])
    ax.axis("off")
    add_panel_label(ax, "A", "Cohort denominator audit")
    draw_box(
        ax,
        (0.02, 0.40),
        0.24,
        0.34,
        "Input/source\ninventory",
        f"{pre_final_total} genomes\n555 rumen\n108 wetland",
        "#D8F3DC",
        body_size=10,
    )
    draw_box(
        ax,
        (0.36, 0.40),
        0.24,
        0.34,
        "Primary POC\ncohort",
        f"{final_total} genomes\n555 rumen\n107 wetland",
        "#BDE0FE",
        body_size=10,
    )
    draw_box(
        ax,
        (0.70, 0.40),
        0.24,
        0.34,
        "Final embedding\nmatrix",
        f"{summary['n_embeddings']} genomes\n{summary['embedding_dim']} dimensions\n0 non-finite",
        "#FFE8D6",
        body_size=10,
    )
    arrow(ax, (0.27, 0.57), (0.35, 0.57))
    arrow(ax, (0.61, 0.57), (0.69, 0.57))
    ax.text(
        0.31,
        0.80,
        f"{excluded} wetland input\nrecord removed",
        ha="center",
        fontsize=9,
        color=COLORS["red"],
        fontweight="bold",
    )
    ax.text(
        0.02,
        0.17,
        "v2.0 correction: keep input and final denominators separate.",
        fontsize=11,
        fontweight="bold",
        color=COLORS["ink"],
    )

    ax = fig.add_subplot(gs[0, 1])
    matrix = pd.crosstab(metadata["ecosystem"], metadata["source"]).reindex(
        index=["rumen", "wetland"], columns=["rumen", "mucc"], fill_value=0
    )
    sns.heatmap(
        matrix,
        annot=True,
        fmt="d",
        cmap="Blues",
        cbar=False,
        linewidths=1.5,
        linecolor="white",
        ax=ax,
        annot_kws={"fontsize": 13, "fontweight": "bold"},
    )
    add_panel_label(ax, "B", "Source/ecosystem confounding")
    ax.set_xlabel("Source project")
    ax.set_ylabel("Ecosystem label")
    ax.text(
        0.5,
        -0.22,
        "All rumen genomes come from PRJEB31266; all wetland genomes come from MUCC.",
        transform=ax.transAxes,
        ha="center",
        fontsize=10,
        color=COLORS["red"],
        fontweight="bold",
    )

    ax = fig.add_subplot(gs[1, 0])
    domain_counts = (
        metadata.groupby(["ecosystem", "domain"]).size().reset_index(name="n")
    )
    domain_pivot = (
        domain_counts.pivot(index="ecosystem", columns="domain", values="n")
        .fillna(0)
        .reindex(index=["rumen", "wetland"])
    )
    domain_pct = domain_pivot.div(domain_pivot.sum(axis=1), axis=0) * 100
    domain_colors = {
        "Archaea": COLORS["archa"],
        "Bacteria": COLORS["bacteria"],
        "Unknown": COLORS["unknown"],
    }
    left = np.zeros(len(domain_pct))
    for domain in ["Archaea", "Bacteria", "Unknown"]:
        vals = domain_pct.get(domain, pd.Series([0, 0], index=domain_pct.index))
        ax.barh(
            domain_pct.index,
            vals,
            left=left,
            color=domain_colors[domain],
            edgecolor="white",
            label=domain,
        )
        left += vals.values
    add_panel_label(ax, "C", "Taxonomic interpretability constraint")
    ax.set_xlabel("Domain label composition (%)")
    ax.set_xlim(0, 100)
    ax.legend(loc="lower right", frameon=True)
    ax.text(
        0.02,
        -0.24,
        "Wetland taxonomy is 100% Unknown in current metadata; mechanism claims require GTDB-Tk/CheckM2/GUNC.",
        transform=ax.transAxes,
        fontsize=10,
        color=COLORS["red"],
        fontweight="bold",
    )

    ax = fig.add_subplot(gs[1, 1])
    ax.axis("off")
    add_panel_label(ax, "D", "Evidence ladder for the current POC")
    rows = [
        ("Proven", "Zero-attrition 662 x 1280 embedding matrix", COLORS["green"]),
        ("Strong", "Ecosystem structure in ESM2 latent space", "#7BDFF2"),
        ("Qualified", "Bridge candidates exist, but mechanism is unassigned", COLORS["gold"]),
        ("Not yet proven", "Methane pathway function, source-robust transfer, MRV readiness", "#F4978E"),
    ]
    y = 0.74
    for label, body, color in rows:
        draw_box(ax, (0.05, y), 0.90, 0.18, label, body, color, 10, 9)
        y -= 0.21

    fig.suptitle(
        "Figure 1. Scientific defensibility begins with denominator and source accounting",
        fontsize=17,
        fontweight="bold",
        color=COLORS["ink"],
    )
    return savefig(fig, "fig_v2_01_evidence_ladder.png")


def make_figure_2(data: dict[str, object]) -> Path:
    """Enhanced embedding geometry and bridge visualization."""
    projection = data["projection"].copy()
    summary = data["summary"]

    top_bridge = projection.sort_values("alpha_transfer_score", ascending=False).head(20)
    bridge_mask = projection["opp_neighbor_fraction"] > 0

    fig = plt.figure(figsize=(15, 12), facecolor=COLORS["paper"])
    gs = gridspec.GridSpec(2, 2, hspace=0.34, wspace=0.28)

    ax = fig.add_subplot(gs[0, 0])
    for eco in ["rumen", "wetland"]:
        m = projection["ecosystem"] == eco
        ax.scatter(
            projection.loc[m, "pc1"],
            projection.loc[m, "pc2"],
            s=24,
            alpha=0.62,
            c=COLORS[eco],
            edgecolors="none",
            label=f"{eco} (n={int(m.sum())})",
        )
    ax.scatter(
        projection.loc[bridge_mask, "pc1"],
        projection.loc[bridge_mask, "pc2"],
        s=110,
        facecolors="none",
        edgecolors=COLORS["gold"],
        linewidths=1.8,
        label="opposite-neighbor candidates",
    )
    add_panel_label(ax, "A", "PCA landscape with bridge candidates")
    ax.set_xlabel("PC1 (44.4%)")
    ax.set_ylabel("PC2 (21.8%)")
    ax.legend(frameon=True, fontsize=8)

    ax = fig.add_subplot(gs[0, 1])
    for eco in ["rumen", "wetland"]:
        m = projection["ecosystem"] == eco
        ax.scatter(
            projection.loc[m, "umap1"],
            projection.loc[m, "umap2"],
            s=24,
            alpha=0.58,
            c=COLORS[eco],
            edgecolors="none",
        )
    ax.scatter(
        top_bridge["umap1"],
        top_bridge["umap2"],
        s=95,
        c=top_bridge["alpha_transfer_score"],
        cmap="inferno",
        marker="*",
        edgecolors="black",
        linewidths=0.45,
        label="top alpha-transfer",
        zorder=10,
    )
    for _, row in top_bridge.head(5).iterrows():
        ax.text(row["umap1"] + 0.08, row["umap2"] + 0.08, safe_short(row["sample"], 13), fontsize=7)
    add_panel_label(ax, "B", "UMAP bridge prioritization")
    ax.set_xlabel("UMAP1")
    ax.set_ylabel("UMAP2")

    ax = fig.add_subplot(gs[1, 0])
    for eco in ["rumen", "wetland"]:
        vals = projection.loc[projection["ecosystem"] == eco, "ecosystem_axis_projection"]
        sns.kdeplot(vals, fill=True, alpha=0.34, color=COLORS[eco], ax=ax, label=eco)
        ax.axvline(vals.mean(), color=COLORS[eco], linewidth=2)
    rumen_vals = projection.loc[
        projection["ecosystem"] == "rumen", "ecosystem_axis_projection"
    ]
    wetland_vals = projection.loc[
        projection["ecosystem"] == "wetland", "ecosystem_axis_projection"
    ]
    ttest = stats.ttest_ind(wetland_vals, rumen_vals, equal_var=False)
    pooled = np.sqrt(
        (
            wetland_vals.var(ddof=1) * (len(wetland_vals) - 1)
            + rumen_vals.var(ddof=1) * (len(rumen_vals) - 1)
        )
        / (len(wetland_vals) + len(rumen_vals) - 2)
    )
    d_val = (wetland_vals.mean() - rumen_vals.mean()) / pooled
    add_panel_label(ax, "C", "Centroid-axis separation")
    ax.set_xlabel("Projection on rumen-to-wetland axis")
    ax.set_ylabel("Density")
    ax.legend(frameon=True)
    ax.text(
        0.02,
        0.95,
        f"Welch t={ttest.statistic:.1f}, p={ttest.pvalue:.1e}; Cohen d={d_val:.2f}",
        transform=ax.transAxes,
        va="top",
        fontsize=10,
        bbox=dict(facecolor="white", edgecolor=COLORS["muted"], alpha=0.9),
    )

    ax = fig.add_subplot(gs[1, 1])
    sns.scatterplot(
        data=projection,
        x="silhouette",
        y="bridge_entropy",
        hue="ecosystem",
        palette={"rumen": COLORS["rumen"], "wetland": COLORS["wetland"]},
        size="opp_neighbor_fraction",
        sizes=(18, 170),
        alpha=0.68,
        edgecolor=None,
        ax=ax,
    )
    ax.axhline(0.5, color=COLORS["gold"], linestyle="--", linewidth=1)
    ax.axvline(0, color=COLORS["muted"], linestyle=":", linewidth=1)
    add_panel_label(ax, "D", "Separation vs. bridge entropy")
    ax.set_xlabel("Per-genome silhouette")
    ax.set_ylabel("Cross-ecosystem neighborhood entropy")
    ax.legend(fontsize=7, frameon=True)
    ax.text(
        0.02,
        0.06,
        f"Global silhouette={summary['silhouette_global']:.3f}; "
        f"bridge candidates={int(bridge_mask.sum())}",
        transform=ax.transAxes,
        fontsize=9,
        bbox=dict(facecolor="white", alpha=0.9, edgecolor="none"),
    )

    fig.suptitle(
        "Figure 2. ESM2 latent geometry is strong, but bridges are sparse and asymmetric",
        fontsize=17,
        fontweight="bold",
        color=COLORS["ink"],
    )
    return savefig(fig, "fig_v2_02_embedding_geometry.png")


def make_figure_3(data: dict[str, object]) -> Path:
    """Bridge candidate ranking and triage."""
    projection = data["projection"].copy()
    top = projection.sort_values("alpha_transfer_score", ascending=False).head(20).copy()
    top["short"] = top["sample"].map(lambda value: safe_short(value, 26))
    top["domain_color"] = top["domain"].map(
        {"Archaea": COLORS["archa"], "Bacteria": COLORS["bacteria"], "Unknown": COLORS["unknown"]}
    ).fillna(COLORS["muted"])

    fig = plt.figure(figsize=(16, 12), facecolor=COLORS["paper"])
    gs = gridspec.GridSpec(2, 2, hspace=0.36, wspace=0.33)

    ax = fig.add_subplot(gs[0, 0])
    top_plot = top.iloc[::-1]
    ax.barh(
        top_plot["short"],
        top_plot["alpha_transfer_score"],
        color=[COLORS[e] for e in top_plot["ecosystem"]],
        edgecolor="white",
    )
    add_panel_label(ax, "A", "Top alpha-transfer candidates")
    ax.set_xlabel("Alpha-transfer score")
    ax.tick_params(axis="y", labelsize=7)
    for y, (_, row) in enumerate(top_plot.iterrows()):
        ax.scatter(
            row["alpha_transfer_score"] + 0.05,
            y,
            s=70,
            color=row["domain_color"],
            edgecolor="black",
            linewidth=0.4,
            zorder=3,
        )
    ax.text(
        0.02,
        -0.16,
        "Bar color = ecosystem; dot color = domain (Archaea/Bacteria/Unknown).",
        transform=ax.transAxes,
        fontsize=9,
        color=COLORS["muted"],
    )

    ax = fig.add_subplot(gs[0, 1])
    sns.scatterplot(
        data=projection,
        x="nearest_same_distance",
        y="nearest_opposite_distance",
        hue="ecosystem",
        palette={"rumen": COLORS["rumen"], "wetland": COLORS["wetland"]},
        size="alpha_transfer_score",
        sizes=(20, 160),
        alpha=0.62,
        edgecolor=None,
        ax=ax,
    )
    ax.plot(
        [projection["nearest_same_distance"].min(), projection["nearest_opposite_distance"].max()],
        [projection["nearest_same_distance"].min(), projection["nearest_opposite_distance"].max()],
        linestyle="--",
        color=COLORS["muted"],
        linewidth=1,
    )
    add_panel_label(ax, "B", "Same-vs-opposite nearest-neighbor distance")
    ax.set_xlabel("Nearest same-ecosystem cosine distance")
    ax.set_ylabel("Nearest opposite-ecosystem cosine distance")
    ax.legend(fontsize=7, frameon=True)

    ax = fig.add_subplot(gs[1, 0])
    ax.axis("off")
    add_panel_label(ax, "C", "Bridge interpretation triage")
    triage = [
        ("Best-supported current signal", "Rumen Archaea with high entropy and nonzero opposite-neighbor fraction"),
        ("Most important unresolved signal", "Wetland candidates remain taxonomically Unknown"),
        ("Mechanistic status", "Unassigned: no MCycDB/KOfam/DRAM2/METABOLIC layer yet"),
        ("Artifact checks still required", "MAG completeness, contamination, chimerism, dereplication"),
    ]
    y = 0.77
    for title, body in triage:
        draw_box(ax, (0.04, y), 0.92, 0.15, title, body, "#E9ECEF", 10, 9)
        y -= 0.19

    ax = fig.add_subplot(gs[1, 1])
    ax.axis("off")
    add_panel_label(ax, "D", "Top bridge candidates requiring mechanism cards")
    cols = ["Rank", "Candidate", "Eco", "Domain", "Opp-NN", "Entropy"]
    table_data = []
    for rank, (_, row) in enumerate(top.head(8).iterrows(), start=1):
        table_data.append(
            [
                rank,
                safe_short(row["sample"], 24),
                row["ecosystem"],
                row["domain"],
                f"{row['opp_neighbor_fraction']:.2f}",
                f"{row['bridge_entropy']:.2f}",
            ]
        )
    table = ax.table(
        cellText=table_data,
        colLabels=cols,
        loc="center",
        cellLoc="left",
        colLoc="left",
        colWidths=[0.08, 0.38, 0.12, 0.13, 0.12, 0.12],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(8)
    table.scale(1, 1.55)
    for (r, c), cell in table.get_celld().items():
        cell.set_edgecolor("#D0D7DE")
        if r == 0:
            cell.set_text_props(weight="bold", color="white")
            cell.set_facecolor(COLORS["blue"])
        elif r % 2 == 0:
            cell.set_facecolor("#F6F8FA")

    fig.suptitle(
        "Figure 3. Bridge candidates are actionable hypotheses, not mechanistic proof yet",
        fontsize=17,
        fontweight="bold",
        color=COLORS["ink"],
    )
    return savefig(fig, "fig_v2_03_bridge_triage.png")


def make_figure_4(data: dict[str, object]) -> Path:
    """Scientific defensibility and missing-evidence audit."""
    fig = plt.figure(figsize=(16, 10), facecolor=COLORS["paper"])
    gs = gridspec.GridSpec(1, 2, width_ratios=[1.15, 0.85], wspace=0.28)

    ax = fig.add_subplot(gs[0, 0])
    rows = [
        "Zero-attrition embedding",
        "Ecosystem latent structure",
        "Bridge candidate detection",
        "Methane mechanism assignment",
        "MAG QC/taxonomy/dereplication",
        "Source-aware transfer validity",
        "MRV-ready feature set",
    ]
    cols = [
        "Current evidence",
        "Independent QC",
        "Functional layer",
        "Source control",
        "Actionable status",
    ]
    values = np.array(
        [
            [2, 1, 0, 0, 2],
            [2, 1, 0, 0, 1],
            [2, 0, 0, 0, 1],
            [0, 0, 0, 0, 0],
            [0, 0, 0, 0, 0],
            [1, 0, 0, 0, 0],
            [0, 0, 0, 0, 0],
        ]
    )
    cmap = matplotlib.colors.ListedColormap(["#F4978E", "#F4D35E", "#2A9D8F"])
    sns.heatmap(
        values,
        cmap=cmap,
        vmin=0,
        vmax=2,
        cbar=False,
        linewidths=1.5,
        linecolor="white",
        xticklabels=cols,
        yticklabels=rows,
        ax=ax,
    )
    add_panel_label(ax, "A", "Claim defensibility matrix")
    ax.tick_params(axis="x", rotation=35, labelsize=9)
    ax.tick_params(axis="y", labelsize=9)
    for i in range(values.shape[0]):
        for j in range(values.shape[1]):
            label = ["Missing", "Partial", "Strong"][values[i, j]]
            ax.text(j + 0.5, i + 0.5, label, ha="center", va="center", fontsize=8)

    ax = fig.add_subplot(gs[0, 1])
    ax.axis("off")
    add_panel_label(ax, "B", "v2.0 critical review: what must change")
    critique = [
        ("1. Strongest result", "A complete, finite 662-genome ESM2 embedding matrix."),
        ("2. Main overclaim risk", "Near-perfect classifier metrics are not independent validation because source and ecosystem are perfectly coupled."),
        ("3. Bridge caveat", "Bridge rank is geometric; methane pathway relevance is not yet measured."),
        ("4. Taxonomy caveat", "All wetland genomes are domain Unknown in current metadata."),
        ("5. MethaNet alignment", "Next value comes from hybrid latent + functional + source-aware features."),
    ]
    y = 0.84
    for title, body in critique:
        draw_box(ax, (0.04, y), 0.92, 0.13, title, body, "#F8F9FA", 9, 8)
        y -= 0.155
    ax.text(
        0.04,
        0.03,
        "Interpretation standard for v2.0: use the POC as a transfer-learning feasibility result, not as a mechanistic methane MRV claim.",
        fontsize=10,
        fontweight="bold",
        color=COLORS["red"],
        wrap=True,
    )

    fig.suptitle(
        "Figure 4. A defensible report must separate signal, caveat, and next evidence",
        fontsize=17,
        fontweight="bold",
        color=COLORS["ink"],
    )
    return savefig(fig, "fig_v2_04_defensibility_audit.png")


def make_figure_5(_: dict[str, object]) -> Path:
    """Functional metagenomics roadmap aligned to MethaNet."""
    fig, ax = plt.subplots(figsize=(17, 9.5), facecolor=COLORS["paper"])
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)

    phases = [
        (
            "Phase A\nMAG QC + identity",
            "CheckM2, GUNC,\nGTDB-Tk, dRep",
            "Every bridge has QC,\ntaxonomy, derep status",
            "#D8F3DC",
        ),
        (
            "Phase B\nMethane mechanism",
            "MCycDB, MCR/Hdr,\nmethylotrophy, AOM,\nsulfur competition",
            "Mechanism cards classify\neach bridge candidate",
            "#BDE0FE",
        ),
        (
            "Phase C\nBroad function",
            "KOfamScan, eggNOG,\nDRAM2, METABOLIC,\ndbCAN, transporters",
            "Annotation missingness\ncannot explain rankings",
            "#FFE8D6",
        ),
        (
            "Phase D\nDeconfounded transfer",
            "k/seed/downsampling,\n2-factor PERMANOVA,\nleave-one-source-out",
            "Signal and bridges survive\nsource-aware controls",
            "#FAD2E1",
        ),
    ]

    ax.text(
        0.5,
        0.965,
        "Figure 5. Functional-metagenomics expansion plan",
        ha="center",
        va="center",
        fontsize=16,
        fontweight="bold",
        color=COLORS["ink"],
    )
    ax.text(
        0.5,
        0.925,
        "from geometric bridges to mechanistic, source-aware MethaNet features",
        ha="center",
        va="center",
        fontsize=12,
        fontweight="bold",
        color=COLORS["muted"],
    )

    x0 = 0.045
    widths = 0.195
    gap = 0.045
    phase_centers = []
    for idx, (title, tools, gate, color) in enumerate(phases):
        x = x0 + idx * (widths + gap)
        phase_centers.append(x + widths / 2)
        draw_box(ax, (x, 0.53), widths, 0.31, title, tools, color, 11, 8)
        arrow(ax, (x + widths / 2, 0.52), (x + widths / 2, 0.43))
        draw_box(
            ax,
            (x, 0.29),
            widths,
            0.14,
            "Gate",
            gate,
            "#FFFFFF",
            9,
            8,
        )
        if idx < len(phases) - 1:
            next_x = x0 + (idx + 1) * (widths + gap)
            arrow(ax, (x + widths + 0.008, 0.685), (next_x - 0.008, 0.685))

    draw_box(
        ax,
        (0.20, 0.055),
        0.60,
        0.13,
        "MethaNet v2.0 endpoint",
        "Hybrid latent + functional + source-aware features\nfor methane MRV prioritization",
        "#E9ECEF",
        12,
        9,
    )
    ax.plot(
        [phase_centers[0], phase_centers[-1]],
        [0.235, 0.235],
        color=COLORS["muted"],
        linewidth=1.3,
    )
    for x in phase_centers:
        ax.plot([x, x], [0.29, 0.235], color=COLORS["muted"], linewidth=1.0)
    arrow(ax, (0.50, 0.235), (0.50, 0.19))
    return savefig(fig, "fig_v2_05_functional_metagenomics_roadmap.png")


def make_figures(data: dict[str, object]) -> list[Path]:
    setup_plotting()
    return [
        make_figure_1(data),
        make_figure_2(data),
        make_figure_3(data),
        make_figure_4(data),
        make_figure_5(data),
    ]


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for name, size, color in [
        ("Title", 24, COLORS["ink"]),
        ("Heading 1", 16, COLORS["blue"]),
        ("Heading 2", 12, COLORS["ink"]),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        style.font.bold = True


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Blue Catalyst Cross-Ecosystem Embedding POC")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(30, 35, 41)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Deep-Dive Validation Report v2.0")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(42, 111, 151)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "662 genomes = 555 rumen + 107 wetland | ESM2-650M latent space | "
        "scientific critique and functional-metagenomics expansion"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(105, 115, 134)
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    p.add_run("\n" + body)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(7.15))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(105, 115, 134)


def add_table_from_rows(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def write_source_manifest() -> None:
    rows = [
        ("Final embeddings", ARTIFACTS / "genome_embeddings.npz"),
        ("Embedding metadata", ARTIFACTS / "embedding_metadata.tsv"),
        ("Embedding stats", ARTIFACTS / "embedding_stats.json"),
        ("Source/input counts", ARTIFACTS / "sample_source_counts.tsv"),
        ("Bridge candidates", ARTIFACTS / "bridging_genomes_top.tsv"),
        ("Projection table", TABLES / "embedding_projection_partial.tsv"),
        ("Analytics summary", ANALYTICS / "analytics_summary.json"),
        ("Original v1 report", SNAPSHOT / "deep_dive_report" / "blue_catalyst_deep_dive_report.docx"),
    ]
    pd.DataFrame(rows, columns=["artifact", "path"]).to_csv(
        SOURCE_MANIFEST,
        sep="\t",
        index=False,
    )


def build_review_markdown(data: dict[str, object]) -> None:
    summary = data["summary"]
    review = f"""# Blue Catalyst Deep-Dive Report v2.0 Review

## Critical conclusion

The current POC is scientifically valuable and worth preserving, but its claim
boundary must be sharper. It proves a complete 662-genome ESM2 embedding and a
strong, ecologically structured latent space. It does not yet prove methane
mechanism, source-independent transfer, or MRV readiness.

## Strong evidence

- Final embedding matrix: {summary['n_embeddings']} x {summary['embedding_dim']}.
- Finite vectors: {summary['n_embeddings']} with zero non-finite vectors.
- Ecosystem counts: {summary['ecosystem_counts']}.
- PERMANOVA: F={summary['permanova']['pseudo_f']:.2f}, p={summary['permanova']['p_value']}, R2={summary['permanova']['r2']:.3f}.
- Silhouette: {summary['silhouette_global']:.3f}.

## Main challenges

1. Source and ecosystem are perfectly confounded.
2. Wetland taxonomy is unresolved in the available metadata.
3. Bridge candidates are latent-space hypotheses, not functionally assigned genomes.
4. Classifier metrics should be framed as separability diagnostics, not validated generalization.
5. Functional-metagenomics layers are missing from the current POC.

## v2.0 response

The v2.0 report adds denominator accounting, source-confounding diagnostics,
bridge triage, a defensibility matrix, and a gated functional-metagenomics
roadmap aligned to MethaNet.
"""
    REVIEW_MD.write_text(review)


def build_docx(data: dict[str, object], figures: list[Path]) -> None:
    projection = data["projection"]
    source_counts = data["source_counts"]
    summary = data["summary"]

    bridge_count = int((projection["opp_neighbor_fraction"] > 0).sum())
    high_entropy_count = int((projection["bridge_entropy"] > 0.5).sum())
    top = projection.sort_values("alpha_transfer_score", ascending=False).head(8)

    doc = Document()
    set_doc_style(doc)
    add_title(doc)

    doc.add_heading("1. Executive Verdict", level=1)
    add_callout(
        doc,
        "v2.0 scientific position",
        (
            "The POC is a strong latent-geometry result: 662 genomes embedded with "
            "zero attrition into a 1,280-dimensional ESM2-650M space, with clear "
            "ecosystem structure and sparse cross-ecosystem bridge candidates. "
            "It is not yet a mechanistic methane MRV result because source, taxonomy, "
            "MAG QC, and functional-metagenomics layers are not resolved."
        ),
    )
    add_bullets(
        doc,
        [
            "Preserve the headline: ESM2 latent space captures ecologically meaningful structure and bridge candidates.",
            "Strengthen the caveat: classifier separability and PERMANOVA are not source-independent validation.",
            "Upgrade the actionability: every bridge candidate must become a QC/taxonomy/function mechanism card.",
            "Align with MethaNet: v2.0 should use the POC as the seed for hybrid latent + functional + deconfounded transfer learning.",
        ],
    )

    add_figure(
        doc,
        figures[0],
        "Figure 1. Denominator integrity, source/ecosystem confounding, taxonomy interpretability, and evidence ladder.",
    )

    doc.add_heading("2. Authoritative Input Datasets", level=1)
    doc.add_paragraph(
        "The v2.0 report is generated from the canonical 662-genome artifact set, "
        "not from the older 40-genome or 400-genome interim reports."
    )
    add_table_from_rows(
        doc,
        ["Input", "Role in v2.0 report"],
        [
            ["genome_embeddings.npz", "Final 662 x 1280 ESM2 embedding matrix"],
            ["embedding_metadata.tsv", "Final source, ecosystem, domain, and protein-count metadata"],
            ["sample_source_counts.tsv", "Pre-final/source denominator: 555 rumen + 108 wetland"],
            ["embedding_stats.json", "Zero-attrition and exclusion accounting"],
            ["embedding_projection_partial.tsv", "PCA, UMAP, t-SNE, bridge, silhouette, and transfer scores"],
            ["bridging_genomes_top.tsv", "Canonical top bridge-candidate table"],
            ["analytics_summary.json", "PERMANOVA, silhouette, CV, caveats, and source-confounding summary"],
        ],
    )

    doc.add_heading("3. What the Current POC Actually Proves", level=1)
    add_bullets(
        doc,
        [
            f"Embedding completeness: {summary['n_embeddings']} genomes with {summary['embedding_dim']} dimensions and no non-finite vectors.",
            f"Final cohort: {summary['ecosystem_counts']['rumen']} rumen and {summary['ecosystem_counts']['wetland']} wetland genomes.",
            f"Pre-final denominator: {int(source_counts['n_samples'].sum())} source-counted genomes, including 108 MUCC/wetland before the final 107 wetland denominator.",
            f"Global structure: PERMANOVA R2={summary['permanova']['r2']:.3f}, F={summary['permanova']['pseudo_f']:.2f}, p={summary['permanova']['p_value']}.",
            f"Bridge structure: {bridge_count} genomes have at least one opposite-ecosystem nearest neighbor; {high_entropy_count} exceed bridge entropy 0.5.",
        ],
    )
    add_figure(
        doc,
        figures[1],
        "Figure 2. Enhanced embedding geometry and bridge distribution across PCA, UMAP, centroid-axis projection, and silhouette/entropy space.",
    )

    doc.add_heading("4. Critical Scientific Challenge", level=1)
    doc.add_paragraph(
        "The original report is compelling, but v2.0 must be more explicit about "
        "where the evidence stops. This improves, rather than weakens, the MethaNet "
        "story: a credible platform is built by converting latent structure into "
        "mechanistically explained and source-aware features."
    )
    add_numbered(
        doc,
        [
            "Classifier AUC/AUPRC of 1.0 should be described as label separability, not validated ecosystem transfer, because source and ecosystem are perfectly coupled.",
            "PERMANOVA confirms structure but does not partition ecosystem from source without multi-source cohorts.",
            "Bridge candidates require mechanism cards; geometry alone cannot tell whether they are methane-relevant, sulfur-associated, substrate-flexible, unclear, or artifacts.",
            "Wetland domain labels are Unknown in the current metadata, limiting biological interpretation until GTDB-Tk and MAG QC are joined.",
            "The current POC has no MCycDB, KO/EC/module, DRAM2, METABOLIC, dbCAN, or transporter layer, so it cannot yet support deep functional metagenomics claims.",
        ],
    )
    add_figure(
        doc,
        figures[3],
        "Figure 4. Claim-level defensibility matrix and concise critical review of the current evidence boundary.",
    )

    doc.add_heading("5. Bridge Candidate Triage", level=1)
    doc.add_paragraph(
        "Bridge genomes are the most important translational output of the POC. "
        "Bridge candidates are actionable hypotheses, not mechanistic proof yet. "
        "In v2.0 they should be handled as prioritized hypotheses: ranked, audited, "
        "assigned mechanism status, and then tested under source-aware controls."
    )
    add_figure(
        doc,
        figures[2],
        "Figure 3. Bridge ranking, neighbor-distance geometry, interpretation triage, and top candidates requiring mechanism cards.",
    )
    add_table_from_rows(
        doc,
        ["Rank", "Candidate", "Ecosystem", "Domain", "Opp-NN fraction", "Entropy", "Alpha score"],
        [
            [
                rank,
                safe_short(row["sample"], 36),
                row["ecosystem"],
                row["domain"],
                f"{row['opp_neighbor_fraction']:.2f}",
                f"{row['bridge_entropy']:.2f}",
                f"{row['alpha_transfer_score']:.2f}",
            ]
            for rank, (_, row) in enumerate(top.iterrows(), start=1)
        ],
    )

    doc.add_heading("6. Functional-Metagenomics v2.0 Roadmap", level=1)
    doc.add_paragraph(
        "The Functional-Metagenomics Expansion Plan for the next outstanding version "
        "of the report should add functional evidence behind explicit gates. The goal "
        "is not more decoration; it is a transition from a strong latent geometry POC "
        "to mechanistically explained, source-aware MethaNet features. Phase D must "
        "include leave-one-source-out validation once additional source projects are "
        "available."
    )
    add_figure(
        doc,
        figures[4],
        "Figure 5. Gated functional-metagenomics expansion plan from MAG QC to source-aware transfer validation.",
    )
    add_table_from_rows(
        doc,
        ["Phase", "Required outputs", "Gate"],
        [
            [
                "A. MAG QC + identity",
                "mag_qc_integrated.tsv with CheckM2, GUNC, GTDB-Tk, dRep status",
                "Every bridge has QC, taxonomy, dereplication status",
            ],
            [
                "B. Methane mechanism",
                "Methane markers, MCycDB hits, pathway completeness, bridge mechanism cards",
                "Top bridges classified as methane-relevant, substrate-flexible, sulfur-associated, unclear, or artifact",
            ],
            [
                "C. Broad function",
                "KO/EC/module matrices, DRAM2/METABOLIC, dbCAN, transporters",
                "Annotation missingness cannot silently explain bridge ranks",
            ],
            [
                "D. Deconfounding",
                "k/seed/downsampling stability, two-factor PERMANOVA, leave-one-source-out validation",
                "Signal and bridge ranking survive source-aware controls",
            ],
        ],
    )

    doc.add_heading("7. v2.0 Report Improvements Implemented", level=1)
    add_bullets(
        doc,
        [
            "Moved caveats from footnote status into the core evidence hierarchy.",
            "Added denominator accounting that reconciles 108 pre-final MUCC/wetland records with the final 107 wetland embeddings.",
            "Reframed bridge candidates as hypotheses requiring mechanism cards.",
            "Added a defensibility matrix to prevent overclaiming while preserving the strength of the POC.",
            "Added an explicit functional-metagenomics roadmap aligned with MethaNet's source-aware transfer-learning objective.",
        ],
    )

    doc.add_heading("8. Methods and Provenance", level=1)
    doc.add_paragraph(
        "Report v2.0 was generated by scripts/build_blue_catalyst_deep_dive_report_v2.py. "
        "All figures are generated from local MethaNet artifacts listed in source_manifest.tsv. "
        "The original report is preserved unmodified."
    )
    doc.add_paragraph(
        f"Generated output path: {REPORT.relative_to(PROJECT)}"
    )

    doc.save(REPORT)


def main() -> None:
    data = load_inputs()
    write_source_manifest()
    build_review_markdown(data)
    figures = make_figures(data)
    build_docx(data, figures)
    print(f"Wrote {REPORT}")
    print(f"Wrote {REVIEW_MD}")
    print(f"Wrote {SOURCE_MANIFEST}")
    for fig in figures:
        print(f"Wrote {fig}")


if __name__ == "__main__":
    main()
